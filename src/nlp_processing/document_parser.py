# -*- coding: utf-8 -*-
"""
文档解析器
支持PDF、Word、Excel、TXT等格式的文档解析
"""

import os
from typing import Dict, List, Optional, Union
from pathlib import Path
import logging

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import openpyxl
except ImportError:
    openpyxl = None

logger = logging.getLogger(__name__)


class DocumentParser:
    """
    文档解析器类
    支持多种文档格式的文本提取
    """
    
    def __init__(self):
        """
        初始化文档解析器
        """
        self.supported_formats = {
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
            '.doc': self._parse_docx,
            '.xlsx': self._parse_xlsx,
            '.xls': self._parse_xlsx,
            '.txt': self._parse_txt,
            '.md': self._parse_txt
        }
    
    def parse_document(self, file_path: Union[str, Path]) -> Dict[str, str]:
        """
        解析文档并提取文本内容
        
        Args:
            file_path: 文档文件路径
            
        Returns:
            包含文档内容的字典，格式为 {'content': '文本内容', 'metadata': {...}}
            
        Raises:
            ValueError: 不支持的文件格式
            FileNotFoundError: 文件不存在
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"不支持的文件格式: {file_extension}")
        
        try:
            parser_func = self.supported_formats[file_extension]
            result = parser_func(file_path)
            
            # 添加基本元数据
            result['metadata'] = result.get('metadata', {})
            result['metadata'].update({
                'file_path': str(file_path),
                'file_name': file_path.name,
                'file_size': file_path.stat().st_size,
                'file_extension': file_extension
            })
            
            return result
            
        except Exception as e:
            logger.error(f"解析文档失败 {file_path}: {str(e)}")
            raise
    
    def _parse_pdf(self, file_path: Path) -> Dict[str, str]:
        """
        解析PDF文档
        
        Args:
            file_path: PDF文件路径
            
        Returns:
            包含文档内容的字典
        """
        if pdfplumber is None:
            raise ImportError("需要安装 pdfplumber 库来解析PDF文件")
        
        content = ""
        metadata = {}
        
        try:
            with pdfplumber.open(file_path) as pdf:
                metadata['page_count'] = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        content += f"\n--- 第{page_num}页 ---\n"
                        content += page_text
                        content += "\n"
                
                # 提取PDF元数据
                if pdf.metadata:
                    metadata.update({
                        'title': pdf.metadata.get('Title', ''),
                        'author': pdf.metadata.get('Author', ''),
                        'subject': pdf.metadata.get('Subject', ''),
                        'creator': pdf.metadata.get('Creator', '')
                    })
        
        except Exception as e:
            logger.error(f"PDF解析错误: {str(e)}")
            raise
        
        return {
            'content': content.strip(),
            'metadata': metadata
        }
    
    def _parse_docx(self, file_path: Path) -> Dict[str, str]:
        """
        解析Word文档
        
        Args:
            file_path: Word文件路径
            
        Returns:
            包含文档内容的字典
        """
        if Document is None:
            raise ImportError("需要安装 python-docx 库来解析Word文件")
        
        content = ""
        metadata = {}
        
        try:
            doc = Document(file_path)
            
            # 提取段落文本
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text.strip())
            
            content = "\n".join(paragraphs)
            
            # 提取表格内容
            tables_content = []
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    table_text.append("\t".join(row_text))
                tables_content.append("\n".join(table_text))
            
            if tables_content:
                content += "\n\n--- 表格内容 ---\n"
                content += "\n\n".join(tables_content)
            
            # 提取文档属性
            core_props = doc.core_properties
            metadata.update({
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else ''
            })
        
        except Exception as e:
            logger.error(f"Word文档解析错误: {str(e)}")
            raise
        
        return {
            'content': content.strip(),
            'metadata': metadata
        }
    
    def _parse_xlsx(self, file_path: Path) -> Dict[str, str]:
        """
        解析Excel文档
        
        Args:
            file_path: Excel文件路径
            
        Returns:
            包含文档内容的字典
        """
        if openpyxl is None:
            raise ImportError("需要安装 openpyxl 库来解析Excel文件")
        
        content = ""
        metadata = {}
        
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            
            sheet_contents = []
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                
                sheet_text = f"\n--- 工作表: {sheet_name} ---\n"
                
                # 获取有数据的区域
                if sheet.max_row > 0 and sheet.max_column > 0:
                    for row in sheet.iter_rows(min_row=1, max_row=sheet.max_row,
                                             min_col=1, max_col=sheet.max_column,
                                             values_only=True):
                        row_text = []
                        for cell_value in row:
                            if cell_value is not None:
                                row_text.append(str(cell_value))
                            else:
                                row_text.append("")
                        
                        if any(cell.strip() for cell in row_text):
                            sheet_text += "\t".join(row_text) + "\n"
                
                sheet_contents.append(sheet_text)
            
            content = "\n".join(sheet_contents)
            
            # 提取工作簿属性
            metadata.update({
                'sheet_count': len(workbook.sheetnames),
                'sheet_names': workbook.sheetnames
            })
            
            # 尝试获取文档属性
            try:
                props = workbook.properties
                metadata.update({
                    'title': props.title or '',
                    'creator': props.creator or '',
                    'created': str(props.created) if props.created else '',
                    'modified': str(props.modified) if props.modified else ''
                })
            except:
                pass
        
        except Exception as e:
            logger.error(f"Excel文档解析错误: {str(e)}")
            raise
        
        return {
            'content': content.strip(),
            'metadata': metadata
        }
    
    def _parse_txt(self, file_path: Path) -> Dict[str, str]:
        """
        解析文本文档
        
        Args:
            file_path: 文本文件路径
            
        Returns:
            包含文档内容的字典
        """
        content = ""
        metadata = {}
        
        try:
            # 尝试不同的编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    metadata['encoding'] = encoding
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise ValueError("无法确定文件编码")
            
            # 统计行数
            metadata['line_count'] = len(content.splitlines())
        
        except Exception as e:
            logger.error(f"文本文档解析错误: {str(e)}")
            raise
        
        return {
            'content': content.strip(),
            'metadata': metadata
        }
    
    def batch_parse(self, file_paths: List[Union[str, Path]]) -> Dict[str, Dict[str, str]]:
        """
        批量解析多个文档
        
        Args:
            file_paths: 文档文件路径列表
            
        Returns:
            字典，键为文件路径，值为解析结果
        """
        results = {}
        
        for file_path in file_paths:
            try:
                result = self.parse_document(file_path)
                results[str(file_path)] = result
            except Exception as e:
                logger.error(f"批量解析失败 {file_path}: {str(e)}")
                results[str(file_path)] = {
                    'content': '',
                    'metadata': {'error': str(e)}
                }
        
        return results
    
    def get_supported_formats(self) -> List[str]:
        """
        获取支持的文件格式列表
        
        Returns:
            支持的文件扩展名列表
        """
        return list(self.supported_formats.keys())
    
    def is_supported_format(self, file_path: Union[str, Path]) -> bool:
        """
        检查文件格式是否支持
        
        Args:
            file_path: 文件路径
            
        Returns:
            是否支持该格式
        """
        file_extension = Path(file_path).suffix.lower()
        return file_extension in self.supported_formats