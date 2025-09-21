# -*- coding: utf-8 -*-
"""
文档解析器
支持PDF、Word、Excel、TXT等格式的文档解析
"""

import os
import logging
from typing import Dict, List, Optional, Any, Union
from abc import ABC, abstractmethod
from pathlib import Path

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import openpyxl
except ImportError:
    openpyxl = None

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None


class DocumentParseError(Exception):
    """文档解析异常"""
    pass


class BaseDocumentParser(ABC):
    """文档解析器基类"""
    
    @abstractmethod
    def can_parse(self, file_path: str) -> bool:
        """检查是否能解析指定文件"""
        pass
    
    @abstractmethod
    def parse(self, file_path: str) -> Dict[str, Any]:
        """解析文档并返回结构化内容"""
        pass


class TxtDocumentParser(BaseDocumentParser):
    """TXT文档解析器"""
    
    def can_parse(self, file_path: str) -> bool:
        """检查是否为TXT文件"""
        return file_path.lower().endswith('.txt')
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """解析TXT文档"""
        try:
            # 尝试多种编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16', 'ascii']
            content = None
            used_encoding = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    used_encoding = encoding
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                raise DocumentParseError(f"无法解码文件 {file_path}")
            
            return {
                'content': content,
                'metadata': {
                    'file_type': 'txt',
                    'encoding': used_encoding,
                    'file_size': os.path.getsize(file_path),
                    'line_count': len(content.splitlines())
                }
            }
            
        except Exception as e:
            raise DocumentParseError(f"解析TXT文档失败: {str(e)}")


class PdfDocumentParser(BaseDocumentParser):
    """PDF文档解析器"""
    
    def can_parse(self, file_path: str) -> bool:
        """检查是否为PDF文件"""
        return file_path.lower().endswith('.pdf')
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """解析PDF文档"""
        if pdfplumber is None and PyPDF2 is None:
            raise DocumentParseError("PDF解析库未安装，请安装 pdfplumber 或 PyPDF2")
        
        try:
            content = ""
            metadata = {
                'file_type': 'pdf',
                'file_size': os.path.getsize(file_path),
                'page_count': 0
            }
            
            # 优先使用pdfplumber
            if pdfplumber:
                content, page_count = self._parse_with_pdfplumber(file_path)
                metadata['page_count'] = page_count
                metadata['parser'] = 'pdfplumber'
            elif PyPDF2:
                content, page_count = self._parse_with_pypdf2(file_path)
                metadata['page_count'] = page_count
                metadata['parser'] = 'PyPDF2'
            
            return {
                'content': content,
                'metadata': metadata
            }
            
        except Exception as e:
            raise DocumentParseError(f"解析PDF文档失败: {str(e)}")
    
    def _parse_with_pdfplumber(self, file_path: str) -> tuple[str, int]:
        """使用pdfplumber解析PDF"""
        content = ""
        page_count = 0
        
        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    content += page_text + "\n"
        
        return content, page_count
    
    def _parse_with_pypdf2(self, file_path: str) -> tuple[str, int]:
        """使用PyPDF2解析PDF"""
        content = ""
        page_count = 0
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            page_count = len(pdf_reader.pages)
            
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    content += page_text + "\n"
        
        return content, page_count


class WordDocumentParser(BaseDocumentParser):
    """Word文档解析器"""
    
    def can_parse(self, file_path: str) -> bool:
        """检查是否为Word文件"""
        return file_path.lower().endswith(('.docx', '.doc'))
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """解析Word文档"""
        if Document is None:
            raise DocumentParseError("Word解析库未安装，请安装 python-docx")
        
        try:
            doc = Document(file_path)
            content = ""
            
            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\n"
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        content += "\t".join(row_text) + "\n"
            
            metadata = {
                'file_type': 'docx',
                'file_size': os.path.getsize(file_path),
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables)
            }
            
            return {
                'content': content,
                'metadata': metadata
            }
            
        except Exception as e:
            raise DocumentParseError(f"解析Word文档失败: {str(e)}")


class ExcelDocumentParser(BaseDocumentParser):
    """Excel文档解析器"""
    
    def can_parse(self, file_path: str) -> bool:
        """检查是否为Excel文件"""
        return file_path.lower().endswith(('.xlsx', '.xls'))
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """解析Excel文档"""
        if openpyxl is None:
            raise DocumentParseError("Excel解析库未安装，请安装 openpyxl")
        
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            content = ""
            sheet_data = {}
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_content = []
                
                for row in sheet.iter_rows(values_only=True):
                    # 过滤空行
                    row_values = [str(cell) if cell is not None else "" for cell in row]
                    if any(value.strip() for value in row_values):
                        sheet_content.append("\t".join(row_values))
                
                if sheet_content:
                    sheet_text = "\n".join(sheet_content)
                    sheet_data[sheet_name] = sheet_text
                    content += f"[工作表: {sheet_name}]\n{sheet_text}\n\n"
            
            metadata = {
                'file_type': 'xlsx',
                'file_size': os.path.getsize(file_path),
                'sheet_count': len(workbook.sheetnames),
                'sheet_names': workbook.sheetnames
            }
            
            return {
                'content': content,
                'metadata': metadata,
                'sheet_data': sheet_data
            }
            
        except Exception as e:
            raise DocumentParseError(f"解析Excel文档失败: {str(e)}")


class DocumentParserFactory:
    """文档解析器工厂类"""
    
    def __init__(self):
        self.parsers = [
            TxtDocumentParser(),
            PdfDocumentParser(),
            WordDocumentParser(),
            ExcelDocumentParser()
        ]
        self.logger = logging.getLogger(__name__)
    
    def get_parser(self, file_path: str) -> Optional[BaseDocumentParser]:
        """根据文件路径获取合适的解析器"""
        for parser in self.parsers:
            if parser.can_parse(file_path):
                return parser
        return None
    
    def parse_document(self, file_path: str) -> Dict[str, Any]:
        """解析文档"""
        if not os.path.exists(file_path):
            raise DocumentParseError(f"文件不存在: {file_path}")
        
        parser = self.get_parser(file_path)
        if parser is None:
            raise DocumentParseError(f"不支持的文件格式: {file_path}")
        
        self.logger.info(f"开始解析文档: {file_path}")
        
        try:
            result = parser.parse(file_path)
            result['file_path'] = file_path
            result['file_name'] = os.path.basename(file_path)
            
            self.logger.info(f"文档解析完成: {file_path}, 内容长度: {len(result.get('content', ''))}")
            return result
            
        except Exception as e:
            self.logger.error(f"文档解析失败: {file_path}, 错误: {str(e)}")
            raise
    
    def parse_documents(self, file_paths: List[str]) -> List[Dict[str, Any]]:
        """批量解析文档"""
        results = []
        errors = []
        
        for file_path in file_paths:
            try:
                result = self.parse_document(file_path)
                results.append(result)
            except Exception as e:
                error_info = {
                    'file_path': file_path,
                    'error': str(e)
                }
                errors.append(error_info)
                self.logger.error(f"解析文档失败: {file_path}, 错误: {str(e)}")
        
        return results, errors
    
    def get_supported_extensions(self) -> List[str]:
        """获取支持的文件扩展名列表"""
        extensions = []
        test_files = {
            '.txt': 'test.txt',
            '.pdf': 'test.pdf',
            '.docx': 'test.docx',
            '.doc': 'test.doc',
            '.xlsx': 'test.xlsx',
            '.xls': 'test.xls'
        }
        
        for ext, test_file in test_files.items():
            parser = self.get_parser(test_file)
            if parser:
                extensions.append(ext)
        
        return extensions


# 全局文档解析器实例
document_parser_factory = DocumentParserFactory()


def parse_document(file_path: str) -> Dict[str, Any]:
    """便捷函数：解析单个文档"""
    return document_parser_factory.parse_document(file_path)


def parse_documents(file_paths: List[str]) -> tuple[List[Dict[str, Any]], List[Dict[str, str]]]:
    """便捷函数：批量解析文档"""
    return document_parser_factory.parse_documents(file_paths)


def get_supported_extensions() -> List[str]:
    """便捷函数：获取支持的文件扩展名"""
    return document_parser_factory.get_supported_extensions()